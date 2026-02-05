import os
import json
import re
import torch
import psutil

from transformers import AutoTokenizer, AutoModelForCausalLM
from peft import PeftModel
from docx import Document

# ===============================
# CONFIG
# ===============================
BASE_MODEL_ID = "microsoft/phi-4-mini-instruct"
ADAPTER_ID = "rmtlabs/phi-4-mini-adapter-v1"

INPUT_DATASET = "dataset_interim_v6.jsonl"   # JSONL or JSON
OUTPUT_DIR = "output_word_cvs"

MAX_NEW_TOKENS = 512

# ===============================
# SYSTEM INFO
# ===============================
def print_system_info():
    print("🖥 System Info")
    print(f"CPU RAM: {psutil.virtual_memory().total / 1e9:.1f} GB")
    print(f"CUDA Available: {torch.cuda.is_available()}")
    if torch.cuda.is_available():
        print(f"GPU: {torch.cuda.get_device_name(0)}")
        print(f"VRAM: {torch.cuda.get_device_properties(0).total_memory / 1e9:.1f} GB")

print_system_info()

torch.backends.cuda.matmul.allow_tf32 = True
torch.backends.cudnn.allow_tf32 = True

# ===============================
# LOAD MODEL
# ===============================
print("Loading Phi-4-Mini base model...")

tokenizer = AutoTokenizer.from_pretrained(BASE_MODEL_ID)

base_model = AutoModelForCausalLM.from_pretrained(
    BASE_MODEL_ID,
    dtype=torch.float16,
    device_map="auto",
    low_cpu_mem_usage=True
)

print("Loading adapter...")

model = PeftModel.from_pretrained(
    base_model,
    ADAPTER_ID,
    dtype=torch.float16
)

model.eval()
print("✅ Phi-4-Mini + Adapter loaded")

# ===============================
# JSON EXTRACTION
# ===============================
def extract_last_json(text: str) -> str:
    matches = re.findall(r"\{[\s\S]*?\}", text)
    if not matches:
        return text

    last = matches[-1]

    # Auto-close JSON if truncated
    open_braces = last.count("{")
    close_braces = last.count("}")
    if open_braces > close_braces:
        last += "}" * (open_braces - close_braces)

    return last

# ===============================
# CV → JSON (MODEL)
# ===============================
def parse_resume(resume_text: str) -> dict:
    messages = [
        {
            "role": "system",
            "content": (
                "You extract structured data from resumes. "
                "Return ONE valid JSON object only. "
                "No markdown. No explanations."
            )
        },
        {
            "role": "user",
            "content": f"""
Convert the following CV into JSON.
Fill this schema using the CV content.
Do not invent information.

{{
  "skills": [],
  "years_of_experience": "",
  "job_roles": [],
  "education": []
}}

CV:
{resume_text}
"""
        }
    ]

    prompt = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True
    )

    inputs = tokenizer(prompt, return_tensors="pt").to(model.device)

    with torch.no_grad():
        output = model.generate(
            **inputs,
            max_new_tokens=MAX_NEW_TOKENS,
            do_sample=False,
            repetition_penalty=1.1,
            eos_token_id=tokenizer.eos_token_id,
            pad_token_id=tokenizer.eos_token_id
        )

    decoded = tokenizer.decode(output[0], skip_special_tokens=True)
    cleaned = extract_last_json(decoded)

    return json.loads(cleaned)

# ===============================
# WRITE WORD DOCUMENT
# ===============================
def write_cv_to_word(cv_json: dict, output_path: str):
    doc = Document()

    doc.add_heading("Curriculum Vitae", level=1)

    if cv_json.get("job_roles"):
        doc.add_heading("Professional Summary", level=2)
        for role in cv_json["job_roles"]:
            doc.add_paragraph(role, style="List Bullet")

    if cv_json.get("years_of_experience"):
        doc.add_heading("Experience", level=2)
        doc.add_paragraph(f"Years of Experience: {cv_json['years_of_experience']}")

    if cv_json.get("skills"):
        doc.add_heading("Skills", level=2)
        for skill in cv_json["skills"]:
            doc.add_paragraph(skill, style="List Bullet")

    if cv_json.get("education"):
        doc.add_heading("Education", level=2)
        for edu in cv_json["education"]:
            line = f"{edu.get('degree', '')} in {edu.get('field', '')} – {edu.get('institution', '')}"
            doc.add_paragraph(line)

    doc.save(output_path)

# ===============================
# DATASET PROCESSING
# ===============================
def process_dataset(dataset_path: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)

    with open(dataset_path, "r") as f:
        for idx, line in enumerate(f):
            record = json.loads(line)

            # Case 1: raw CV text
            if "cv_text" in record:
                cv_json = parse_resume(record["cv_text"])
            else:
                # Case 2: already structured JSON
                cv_json = record

            output_path = os.path.join(output_dir, f"cv_{idx + 1}.docx")
            write_cv_to_word(cv_json, output_path)

            print(f"✅ Generated {output_path}")

# ===============================
# MAIN
# ===============================
if __name__ == "__main__":
    process_dataset(INPUT_DATASET, OUTPUT_DIR)
